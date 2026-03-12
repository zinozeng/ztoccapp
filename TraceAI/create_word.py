#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    title = doc.add_heading('鲸小宝重构立项文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    sections = [
        ('一、项目基本信息', add_basic_info),
        ('二、项目背景与目标', add_background),
        ('三、项目范围', add_scope),
        ('四、技术方案', add_tech),
        ('五、项目计划', add_plan),
        ('六、资源需求', add_resources),
        ('七、风险评估', add_risks),
        ('八、质量保障', add_quality),
        ('九、成本预算', add_budget),
        ('十、成功标准', add_success),
        ('十一、项目组织结构', add_org),
        ('十二、附录', add_appendix),
    ]
    
    for section_title, section_func in sections:
        section_func(doc)
    
    add_signatures(doc)
    
    doc.save('鲸小宝重构立项文档.docx')
    print('✓ Word 文档已生成：鲸小宝重构立项文档.docx')

def add_basic_info(doc):
    doc.add_heading('一、项目基本信息', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('项目名称', '鲸小宝重构项目'),
        ('项目负责人', '待分配'),
        ('项目起止时间', '2025 年 1 月 - 2025 年 6 月'),
        ('项目预算', '待评估'),
        ('优先级', 'P0'),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

def add_background(doc):
    doc.add_heading('二、项目背景与目标', level=1)
    doc.add_heading('2.1 项目背景', level=2)
    p = doc.add_paragraph()
    p.add_run('鲸小宝作为核心业务系统，当前架构存在以下问题：')
    for item in ['技术栈老旧，维护成本高', '系统性能瓶颈明显', '扩展性差，难以支撑业务快速发展', '代码质量参差不齐，技术债务累积']:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('2.2 项目目标', level=2)
    for i, goal in enumerate(['技术架构升级：采用现代化技术栈，提升系统可维护性', '性能优化：核心接口响应时间降低 50%', '可扩展性提升：支持业务快速迭代和横向扩展', '代码质量提升：建立完善的代码规范和测试体系'], 1):
        doc.add_paragraph(f'{i}. {goal}')
    doc.add_paragraph()

def add_scope(doc):
    doc.add_heading('三、项目范围', level=1)
    doc.add_heading('3.1 包含内容', level=2)
    for item in ['前端架构重构', '后端服务拆分与优化', '数据库架构升级', '中间件选型与迁移', 'DevOps 体系建设']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('3.2 不包含内容', level=2)
    for item in ['业务流程重构', '第三方系统接口变更']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_tech(doc):
    doc.add_heading('四、技术方案', level=1)
    doc.add_heading('4.1 技术栈选型', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    data = [
        ('层级', '原技术栈', '新技术栈'),
        ('前端框架', 'Vue 2.x', 'Vue 3.x + TypeScript'),
        ('状态管理', 'Vuex', 'Pinia'),
        ('构建工具', 'Webpack 4', 'Vite 5'),
        ('后端框架', 'Spring Boot 2.x', 'Spring Boot 3.x'),
        ('数据库', 'MySQL 5.7', 'MySQL 8.0'),
        ('缓存', 'Redis 3.x', 'Redis 7.x'),
        ('消息队列', 'RabbitMQ', 'RocketMQ'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('4.2 架构设计', level=2)
    doc.add_heading('4.2.1 前端架构', level=3)
    for item in ['组件化开发', '微前端架构', '统一 UI 组件库']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('4.2.2 后端架构', level=3)
    for item in ['领域驱动设计 (DDD)', '服务拆分', 'API 网关统一接入']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('4.2.3 数据架构', level=3)
    for item in ['读写分离', '分库分表', '数据迁移方案']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_plan(doc):
    doc.add_heading('五、项目计划', level=1)
    doc.add_heading('5.1 阶段划分', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    data = [
        ('阶段', '时间', '主要任务', '交付物'),
        ('第一阶段', '2025.01-2025.02', '需求分析、技术方案设计', '需求文档、技术方案文档'),
        ('第二阶段', '2025.02-2025.03', '基础设施搭建、核心模块开发', '基础框架、核心模块'),
        ('第三阶段', '2025.03-2025.04', '业务模块开发、数据迁移', '业务模块、迁移方案'),
        ('第四阶段', '2025.04-2025.05', '系统集成测试、性能优化', '测试报告、性能报告'),
        ('第五阶段', '2025.05-2025.06', '灰度发布、正式上线', '上线报告、运维文档'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('5.2 里程碑', level=2)
    for i, m in enumerate(['M1 (2025.02.28): 完成技术方案评审', 'M2 (2025.03.31): 完成核心模块开发', 'M3 (2025.04.30): 完成所有业务模块开发', 'M4 (2025.05.31): 完成系统测试', 'M5 (2025.06.30): 正式上线'], 1):
        doc.add_paragraph(m)
    doc.add_paragraph()

def add_resources(doc):
    doc.add_heading('六、资源需求', level=1)
    doc.add_heading('6.1 人力资源', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    data = [
        ('角色', '人数', '职责'),
        ('项目经理', '1', '项目整体管理'),
        ('架构师', '1', '技术架构设计'),
        ('前端开发', '3', '前端开发'),
        ('后端开发', '5', '后端开发'),
        ('测试工程师', '2', '测试工作'),
        ('UI 设计师', '1', '界面设计'),
        ('运维工程师', '1', '部署运维'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('6.2 硬件资源', level=2)
    for item in ['开发环境服务器：4 台', '测试环境服务器：4 台', '生产环境服务器：8 台', '数据库服务器：4 台']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('6.3 软件资源', level=2)
    for item in ['开发工具授权', '第三方服务采购', '云资源费用']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_risks(doc):
    doc.add_heading('七、风险评估', level=1)
    
    doc.add_heading('7.1 技术风险', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('技术选型不当', '中', '高', '充分调研，POC 验证'),
        ('数据迁移失败', '中', '高', '制定详细迁移方案，多次演练'),
        ('性能不达标', '低', '高', '早期性能测试，持续优化'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('7.2 项目风险', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('人员流失', '中', '中', '知识共享，文档沉淀'),
        ('需求变更', '高', '中', '严格控制需求范围'),
        ('进度延期', '中', '中', '敏捷开发，及时调整'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('7.3 业务风险', level=2)
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('业务中断', '低', '高', '灰度发布，快速回滚'),
        ('数据丢失', '低', '高', '完整备份，数据校验'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

def add_quality(doc):
    doc.add_heading('八、质量保障', level=1)
    doc.add_heading('8.1 代码质量', level=2)
    for item in ['代码审查制度', '单元测试覆盖率>80%', '集成测试全覆盖', '代码静态分析']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('8.2 测试策略', level=2)
    for item in ['单元测试', '接口测试', '性能测试', '安全测试', '兼容性测试']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('8.3 监控体系', level=2)
    for item in ['应用监控', '日志监控', '性能监控', '业务监控']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_budget(doc):
    doc.add_heading('九、成本预算', level=1)
    doc.add_heading('9.1 人力成本', level=2)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    data = [
        ('角色', '人月', '单价 (元)', '小计 (元)'),
        ('项目经理', '6', '40000', '240000'),
        ('架构师', '6', '45000', '270000'),
        ('前端开发', '18', '30000', '540000'),
        ('后端开发', '30', '30000', '900000'),
        ('测试工程师', '12', '25000', '300000'),
        ('UI 设计师', '3', '28000', '84000'),
        ('运维工程师', '6', '28000', '168000'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    total_row = table.add_row()
    total_row.cells[0].text = '合计'
    total_row.cells[3].text = '2,502,000'
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True
    total_row.cells[3].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('9.2 硬件成本', level=2)
    for item in ['服务器采购：500,000 元', '网络设备：100,000 元']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('9.3 软件成本', level=2)
    for item in ['软件授权：200,000 元', '云服务：300,000 元']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('9.4 总预算', level=2)
    p = doc.add_paragraph()
    run = p.add_run('总计：3,602,000 元')
    run.font.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

def add_success(doc):
    doc.add_heading('十、成功标准', level=1)
    doc.add_heading('10.1 技术指标', level=2)
    for item in ['核心接口响应时间 < 200ms', '系统可用性 > 99.9%', '并发支持能力提升 3 倍', '页面加载时间 < 2s']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('10.2 业务指标', level=2)
    for item in ['用户满意度提升 30%', '系统故障率降低 80%', '新功能上线周期缩短 50%']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('10.3 质量指标', level=2)
    for item in ['线上严重 Bug 数为 0', '代码评审通过率 100%', '测试用例通过率 100%']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_org(doc):
    doc.add_heading('十一、项目组织结构', level=1)
    doc.add_heading('11.1 组织架构', level=2)
    org_text = '''项目指导委员会
    |
项目经理
    |
    +-- 技术负责人
    |       |-- 前端组
    |       |-- 后端组
    |       |-- 测试组
    |
    +-- 产品负责人
    |       |-- 需求分析
    |       |-- 用户体验
    |
    +-- 运维负责人
            |-- 部署运维
            |-- 监控告警'''
    doc.add_paragraph(org_text)
    doc.add_heading('11.2 沟通机制', level=2)
    for item in ['每日站会：每天 9:30，15 分钟', '周例会：每周一，1 小时', '月度汇报：每月最后一个工作日', '技术评审：按需召开']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

def add_appendix(doc):
    doc.add_heading('十二、附录', level=1)
    doc.add_heading('12.1 参考文档', level=2)
    for item in ['鲸小宝现状分析报告', '技术选型调研报告', '竞品分析报告']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('12.2 术语表', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    data = [
        ('术语', '说明'),
        ('DDD', '领域驱动设计'),
        ('POC', '概念验证'),
        ('DevOps', '开发运维一体化'),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

def add_signatures(doc):
    doc.add_heading('审批签字', level=1)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    data = [
        ('角色', '姓名', '签字', '日期'),
        ('项目经理', '', '', ''),
        ('技术负责人', '', '', ''),
        ('产品负责人', '', '', ''),
        ('部门总监', '', '', ''),
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            sign_table = table
            sign_table.rows[i].cells[j].text = cell
            if i == 0:
                sign_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('文档版本：v1.0    创建日期：2025 年 1 月    最后更新：2025 年 1 月')
    footer_run.font.size = Pt(10)

if __name__ == '__main__':
    create_word_document()
