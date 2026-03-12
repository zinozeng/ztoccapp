#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('鲸小宝重构立项文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # 一、项目背景
    doc.add_heading('一、项目背景', level=1)
    p = doc.add_paragraph()
    p.add_run('鲸小宝重构项目是基于对市场现状、竞争态势的深度分析，以及公司业务发展的战略需求而提出的重要项目。').font.bold = True
    
    doc.add_heading('1.1 市场现状', level=2)
    market = [
        '冷链物流行业数字化转型加速',
        '客户对物流可视化、智能化需求日益增长',
        '竞争对手纷纷推出智能化产品',
        '传统人工管理模式已无法满足业务发展需要',
    ]
    for item in market:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('1.2 现有系统问题', level=2)
    problems = [
        '技术栈老旧：Vue 2.x、Spring Boot 2.x 等技术版本已过时',
        '系统性能瓶颈：核心接口响应时间长，并发能力不足',
        '扩展性差：难以支撑业务快速迭代和横向扩展',
        '代码质量参差不齐：技术债务累积，维护成本高',
        '用户体验待提升：界面交互不够流畅，功能操作复杂',
    ]
    for item in problems:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('1.3 重构必要性', level=2)
    necessity = doc.add_paragraph()
    necessity.add_run('通过全面重构，打造一款具有竞争力的智能化冷链物流管理系统，提升客户满意度，降低运营成本，为公司创造更大价值。').font.italic = True
    
    doc.add_paragraph()
    
    # 二、总体要求
    doc.add_heading('二、总体要求', level=1)
    
    doc.add_heading('2.1 建设目标', level=2)
    goals = doc.add_paragraph()
    goals.add_run('打造一款行业领先的智能化冷链物流管理平台，实现以下核心目标：').font.bold = True
    
    target_table = doc.add_table(rows=5, cols=2)
    target_table.style = 'Table Grid'
    target_data = [
        ('目标维度', '具体描述'),
        ('技术先进性', '采用业界主流的技术栈和架构设计，确保系统 3-5 年不落后'),
        ('业务支撑能力', '全面支撑公司冷链物流业务，支持未来业务规模 3 倍增长'),
        ('用户体验', '打造极致的用户体验，操作简便，界面友好'),
        ('运维效率', '建立完善的监控运维体系，降低运维成本 50%'),
    ]
    for i, row in enumerate(target_data):
        for j, cell in enumerate(row):
            target_table.rows[i].cells[j].text = cell
            if i == 0:
                target_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 设计原则', level=2)
    principles = [
        '先进性原则：采用业界成熟且先进的技术架构',
        '可靠性原则：系统稳定可靠，核心业务可用性达到 99.9%',
        '扩展性原则：支持水平扩展，满足业务增长需求',
        '安全性原则：建立完善的安全防护体系，保障数据安全',
        '易用性原则：界面简洁直观，操作便捷',
    ]
    for item in principles:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_paragraph()
    
    # 三、人员方面
    doc.add_heading('三、人员方面', level=1)
    
    doc.add_heading('3.1 团队组建要求', level=2)
    team = [
        '前端开发人员：精通 Vue 3.x、TypeScript，有大型项目经验',
        '后端开发人员：熟练掌握 Spring Boot、微服务架构',
        '架构师：10 年以上经验，有类似项目成功案例',
        '产品经理：熟悉冷链物流业务，有 B 端产品设计经验',
        '测试工程师：精通自动化测试，有性能测试经验',
    ]
    for item in team:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.2 团队规模', level=2)
    doc.add_paragraph('建议组建 15-20 人的专职研发团队，确保项目按时高质量交付。')
    
    doc.add_paragraph()
    
    # 四、总体规划方面
    doc.add_heading('四、总体规划方面', level=1)
    
    doc.add_heading('4.1 总体架构', level=2)
    arch = doc.add_paragraph()
    run = arch.add_run('采用微服务架构，将系统拆分为多个独立的服务模块，包括：用户中心、订单管理、运单管理、仓储管理、运输管理、结算中心等核心模块。')
    run.font.size = Pt(11)
    
    doc.add_heading('4.2 技术路线', level=2)
    tech_route = [
        '前端：Vue 3.x + TypeScript + Vite + Element Plus',
        '后端：Spring Boot 3.x + Spring Cloud Alibaba',
        '数据库：MySQL 8.0（主）+ Redis 7.x（缓存）',
        '消息队列：RocketMQ',
        '容器化：Docker + Kubernetes',
        ' DevOps：Jenkins + GitLab CI/CD',
    ]
    for item in tech_route:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.3 实施策略', level=2)
    strategy = doc.add_paragraph()
    strategy.add_run('采用"整体规划、分步实施、渐进式重构"的策略，确保业务连续性的同时完成系统升级。').font.italic = True
    
    doc.add_paragraph()
    
    # 五、功能方面
    doc.add_heading('五、功能方面', level=1)
    
    doc.add_heading('5.1 功能优先级划分', level=2)
    
    # P0 优先级
    doc.add_heading('5.1.1 P0 - 核心功能（必须实现）', level=3)
    p0_table = doc.add_table(rows=5, cols=2)
    p0_table.style = 'Table Grid'
    p0_data = [
        ('功能模块', '功能描述'),
        ('运单管理', '支持运单的创建、查询、修改、删除，支持批量操作和导入导出'),
        ('异常运单管理', '支持异常运单的登记、处理、跟踪，支持异常类型分类和统计分析'),
        ('客户管理', '客户信息管理、合同管理、信用管理，支持客户分级和权限控制'),
        ('基础数据管理', '网点管理、车辆管理、人员管理、线路管理等基础数据维护'),
    ]
    for i, row in enumerate(p0_data):
        for j, cell in enumerate(row):
            p0_table.rows[i].cells[j].text = cell
            if i == 0:
                p0_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    # P1 优先级
    doc.add_heading('5.1.2 P1 - 重要功能（应该实现）', level=3)
    p1_table = doc.add_table(rows=5, cols=2)
    p1_table.style = 'Table Grid'
    p1_data = [
        ('功能模块', '功能描述'),
        ('智能调度', '基于 AI 算法的智能路径规划、车辆调度，优化运输效率'),
        ('温度监控', '实时温度监控、异常预警、温度曲线展示，支持设备接入'),
        ('报表统计', '多维度业务报表、数据可视化大屏、自定义报表配置'),
        ('移动应用', '司机 APP、业务员 APP、管理端 APP，支持移动端办公'),
    ]
    for i, row in enumerate(p1_data):
        for j, cell in enumerate(row):
            p1_table.rows[i].cells[j].text = cell
            if i == 0:
                p1_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    # P2 优先级
    doc.add_heading('5.1.3 P2 - 优化功能（可以实现）', level=3)
    p2_table = doc.add_table(rows=4, cols=2)
    p2_table.style = 'Table Grid'
    p2_data = [
        ('功能模块', '功能描述'),
        ('AI 预测', '基于历史数据的货量预测、运力预测、成本预测'),
        ('区块链溯源', '基于区块链的冷链溯源，提升产品可信度'),
        ('物联网集成', '与 IoT 设备深度集成，实现设备自动化管理'),
    ]
    for i, row in enumerate(p2_data):
        for j, cell in enumerate(row):
            p2_table.rows[i].cells[j].text = cell
            if i == 0:
                p2_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('5.2 其他需求', level=3)
    other = doc.add_paragraph()
    other.add_run('系统应支持多租户、多语言、多币种，满足集团化管控和国际化业务拓展需求。')
    
    doc.add_paragraph()
    
    # 六、具体的产品需求
    doc.add_heading('六、具体的产品需求', level=1)
    
    doc.add_heading('6.1 运单管理模块', level=2)
    waybill = [
        '支持运单快速创建：扫码录入、模板导入、API 对接',
        '运单状态实时跟踪：从发货到签收全流程可视化',
        '电子签收：支持电子签名、拍照签收',
        '运单打印：支持自定义模板批量打印',
        '异常处理：异常自动识别、智能分派、处理跟踪',
    ]
    for item in waybill:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.2 温度监控模块', level=2)
    temp = [
        '实时温度采集：支持多种温度设备接入，采集频率可配置',
        '温度预警：超温自动报警，支持短信、APP 推送、邮件多种通知方式',
        '温度曲线：可视化展示温度变化趋势，支持导出',
        '温度报表：按车次、客户、时间等多维度统计',
    ]
    for item in temp:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.3 智能调度模块', level=2)
    dispatch = [
        '智能配载：根据货物体积、重量、温度要求自动配载',
        '路径优化：基于路况、时效、成本的多目标路径优化',
        '车辆调度：自动匹配可用车辆，支持人工调整',
        '司机管理：司机排班、绩效考核、培训管理',
    ]
    for item in dispatch:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.4 报表统计模块', level=2)
    report = [
        '运营报表：货量统计、收入统计、成本分析',
        '质量报表：准时率、破损率、投诉率分析',
        '客户分析：客户贡献度、客户流失预警',
        '自定义报表：支持用户自定义报表模板',
        '数据大屏：关键指标实时展示，支持大屏投放',
    ]
    for item in report:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # 七、项目计划
    doc.add_heading('七、项目计划', level=1)
    
    doc.add_heading('7.1 阶段划分', level=2)
    plan_table = doc.add_table(rows=6, cols=4)
    plan_table.style = 'Table Grid'
    plan_data = [
        ('阶段', '时间', '主要任务', '交付物'),
        ('第一阶段', '2025.01-2025.02', '需求分析、技术方案设计', '需求文档、技术方案文档'),
        ('第二阶段', '2025.02-2025.03', '基础设施搭建、核心模块开发', '基础框架、核心模块'),
        ('第三阶段', '2025.03-2025.04', '业务模块开发、数据迁移', '业务模块、迁移方案'),
        ('第四阶段', '2025.04-2025.05', '系统集成测试、性能优化', '测试报告、性能报告'),
        ('第五阶段', '2025.05-2025.06', '灰度发布、正式上线', '上线报告、运维文档'),
    ]
    for i, row in enumerate(plan_data):
        for j, cell in enumerate(row):
            plan_table.rows[i].cells[j].text = cell
            if i == 0:
                plan_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('7.2 里程碑', level=2)
    milestones = [
        'M1 (2025.02.28): 完成需求调研和技术方案评审',
        'M2 (2025.03.31): 完成 P0 核心功能开发',
        'M3 (2025.04.30): 完成 P1 重要功能开发',
        'M4 (2025.05.31): 完成系统测试和用户培训',
        'M5 (2025.06.30): 正式上线运行',
    ]
    for i, m in enumerate(milestones, 1):
        doc.add_paragraph(f'{i}. {m}')
    
    doc.add_paragraph()
    
    # 八、资源需求
    doc.add_heading('八、资源需求', level=1)
    
    doc.add_heading('8.1 人力资源', level=2)
    hr_table = doc.add_table(rows=8, cols=3)
    hr_table.style = 'Table Grid'
    hr_data = [
        ('角色', '人数', '职责'),
        ('项目经理', '1', '项目整体管理'),
        ('架构师', '1', '技术架构设计'),
        ('前端开发', '3', '前端开发'),
        ('后端开发', '5', '后端开发'),
        ('测试工程师', '2', '测试工作'),
        ('UI 设计师', '1', '界面设计'),
        ('运维工程师', '1', '部署运维'),
    ]
    for i, row in enumerate(hr_data):
        for j, cell in enumerate(row):
            hr_table.rows[i].cells[j].text = cell
            if i == 0:
                hr_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('8.2 硬件资源', level=2)
    hardware = [
        '开发环境服务器：4 台',
        '测试环境服务器：4 台',
        '生产环境服务器：8 台',
        '数据库服务器：4 台',
    ]
    for item in hardware:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # 九、风险评估
    doc.add_heading('九、风险评估', level=1)
    
    doc.add_heading('9.1 技术风险', level=2)
    tech_risk_table = doc.add_table(rows=4, cols=4)
    tech_risk_table.style = 'Table Grid'
    tech_risk_data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('技术选型不当', '中', '高', '充分调研，POC 验证'),
        ('数据迁移失败', '中', '高', '制定详细迁移方案，多次演练'),
        ('性能不达标', '低', '高', '早期性能测试，持续优化'),
    ]
    for i, row in enumerate(tech_risk_data):
        for j, cell in enumerate(row):
            tech_risk_table.rows[i].cells[j].text = cell
            if i == 0:
                tech_risk_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('9.2 项目风险', level=2)
    proj_risk_table = doc.add_table(rows=4, cols=4)
    proj_risk_table.style = 'Table Grid'
    proj_risk_data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('人员流失', '中', '中', '知识共享，文档沉淀'),
        ('需求变更', '高', '中', '严格控制需求范围'),
        ('进度延期', '中', '中', '敏捷开发，及时调整'),
    ]
    for i, row in enumerate(proj_risk_data):
        for j, cell in enumerate(row):
            proj_risk_table.rows[i].cells[j].text = cell
            if i == 0:
                proj_risk_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('9.3 业务风险', level=2)
    biz_risk_table = doc.add_table(rows=3, cols=4)
    biz_risk_table.style = 'Table Grid'
    biz_risk_data = [
        ('风险项', '可能性', '影响程度', '应对措施'),
        ('业务中断', '低', '高', '灰度发布，快速回滚'),
        ('数据丢失', '低', '高', '完整备份，数据校验'),
    ]
    for i, row in enumerate(biz_risk_data):
        for j, cell in enumerate(row):
            biz_risk_table.rows[i].cells[j].text = cell
            if i == 0:
                biz_risk_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    # 十、成本预算
    doc.add_heading('十、成本预算', level=1)
    
    doc.add_heading('10.1 人力成本', level=2)
    cost_table = doc.add_table(rows=8, cols=4)
    cost_table.style = 'Table Grid'
    cost_data = [
        ('角色', '人月', '单价 (元)', '小计 (元)'),
        ('项目经理', '6', '40000', '240000'),
        ('架构师', '6', '45000', '270000'),
        ('前端开发', '18', '30000', '540000'),
        ('后端开发', '30', '30000', '900000'),
        ('测试工程师', '12', '25000', '300000'),
        ('UI 设计师', '3', '28000', '84000'),
        ('运维工程师', '6', '28000', '168000'),
    ]
    for i, row in enumerate(cost_data):
        for j, cell in enumerate(row):
            cost_table.rows[i].cells[j].text = cell
            if i == 0:
                cost_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    total_row = cost_table.add_row()
    total_row.cells[0].text = '合计'
    total_row.cells[3].text = '2,502,000'
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True
    total_row.cells[3].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    
    doc.add_heading('10.2 硬件成本', level=2)
    hardware_cost = [
        '服务器采购：500,000 元',
        '网络设备：100,000 元',
    ]
    for item in hardware_cost:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.3 软件成本', level=2)
    software_cost = [
        '软件授权：200,000 元',
        '云服务：300,000 元',
    ]
    for item in software_cost:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('10.4 总预算', level=2)
    total = doc.add_paragraph()
    run = total.add_run('总计：3,602,000 元')
    run.font.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    
    # 十一、成功标准
    doc.add_heading('十一、成功标准', level=1)
    
    doc.add_heading('11.1 技术指标', level=2)
    tech_success = [
        '核心接口响应时间 < 200ms',
        '系统可用性 > 99.9%',
        '并发支持能力提升 3 倍',
        '页面加载时间 < 2s',
    ]
    for item in tech_success:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('11.2 业务指标', level=2)
    biz_success = [
        '用户满意度提升 30%',
        '系统故障率降低 80%',
        '新功能上线周期缩短 50%',
    ]
    for item in biz_success:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('11.3 质量指标', level=2)
    quality_success = [
        '线上严重 Bug 数为 0',
        '代码评审通过率 100%',
        '测试用例通过率 100%',
    ]
    for item in quality_success:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # 十二、待办事项
    doc.add_heading('十二、待办事项', level=1)
    
    todo = [
        '确定项目组织架构和人员配置',
        '完成详细的需求调研和分析',
        '组织技术方案评审',
        '制定详细的项目计划',
        '准备开发、测试环境',
        '启动供应商选型和采购流程',
    ]
    for i, item in enumerate(todo, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 审批签字
    doc.add_heading('审批签字', level=1)
    sign_table = doc.add_table(rows=5, cols=4)
    sign_table.style = 'Table Grid'
    sign_data = [
        ('角色', '姓名', '签字', '日期'),
        ('项目经理', '', '', ''),
        ('技术负责人', '', '', ''),
        ('产品负责人', '', '', ''),
        ('部门总监', '', '', ''),
    ]
    for i, row in enumerate(sign_data):
        for j, cell in enumerate(row):
            sign_table.rows[i].cells[j].text = cell
            if i == 0:
                sign_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 版本信息
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('文档版本：v2.0    创建日期：2025 年 1 月    最后更新：2025 年 1 月')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save('鲸小宝重构立项文档.docx')
    print('✓ Word 文档已更新：鲸小宝重构立项文档.docx')

if __name__ == '__main__':
    create_word_document()
